"""
12_export_docx.py
Convert docs/manuscript_v1.md → docs/manuscript_v1.docx
Applies Elsevier-style formatting (Times New Roman 12pt, double-spaced,
numbered headings, caption style, table style).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from pathlib import Path

ROOT   = Path(__file__).resolve().parent.parent
MD     = ROOT / "docs" / "manuscript_v1.md"
DOCX   = ROOT / "docs" / "manuscript_v1.docx"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def set_para_spacing(para, before=0, after=6, line=24):
    """line in pt*20 (twips). 24 = double-space at 12pt."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)

def add_heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 12}
    bold  = {1: True, 2: True, 3: True}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=bold.get(level, False))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    # Handle inline bold **text** and italic *text*
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=10)
        else:
            run = p.add_run(part)
            set_run_font(run)
    set_para_spacing(p, before=0, after=6, line=24)
    p.paragraph_format.first_line_indent = Cm(0.75)
    return p

def add_formula(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    set_run_font(run, name="Times New Roman", size=12, italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=3, after=3, line=18)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)
    set_para_spacing(p, before=2, after=8, line=15)
    return p

def add_table_from_md(doc, rows):
    """rows: list of list of strings (already stripped of | chars)"""
    if len(rows) < 2:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        if all(set(cell.strip()) <= {'-', ':', ' '} for cell in row):
            continue  # skip separator row
        tr = table.add_row()
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            bold = (r_idx == 0)
            set_run_font(run, size=10, bold=bold)
            p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()  # space after table

# ── parse and render ─────────────────────────────────────────────────────────

doc = Document()

# Page margins (Elsevier-style: 2.5 cm all round)
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Title
lines = MD.read_text(encoding="utf-8").splitlines()

title_done = False
abstract_mode = False
table_buf  = []
in_table   = False

i = 0
while i < len(lines):
    line = lines[i]

    # ── document title (first # heading) ──────────────────────────────────
    if line.startswith('# ') and not title_done:
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        set_run_font(run, size=18, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        title_done = True
        i += 1
        continue

    # ── journal/backup meta lines (bold italic) ───────────────────────────
    if line.startswith('**Journal:**') or line.startswith('**Backup:**'):
        p = doc.add_paragraph()
        run = p.add_run(line.replace('**', ''))
        set_run_font(run, size=10, italic=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # ── blank / separator ─────────────────────────────────────────────────
    if line.strip() == '' or line.strip() == '---':
        if in_table and table_buf:
            add_table_from_md(doc, table_buf)
            table_buf = []
            in_table  = False
        i += 1
        continue

    # ── headings ──────────────────────────────────────────────────────────
    if line.startswith('### '):
        add_heading(doc, line[4:], 3)
        i += 1
        continue
    if line.startswith('## '):
        add_heading(doc, line[3:], 2)
        i += 1
        continue
    if line.startswith('# '):
        add_heading(doc, line[2:], 1)
        i += 1
        continue

    # ── tables ────────────────────────────────────────────────────────────
    if line.startswith('|'):
        in_table = True
        cols = [c for c in line.split('|') if c != '']
        table_buf.append(cols)
        i += 1
        continue
    else:
        if in_table and table_buf:
            add_table_from_md(doc, table_buf)
            table_buf = []
            in_table  = False

    # ── blockquotes / formulas (lines starting with >) ───────────────────
    if line.startswith('> '):
        add_formula(doc, line[2:])
        i += 1
        continue

    # ── bullet lists ──────────────────────────────────────────────────────
    if line.startswith('- ') or re.match(r'^\d+\. ', line):
        p = doc.add_paragraph()
        text = re.sub(r'^[-\d]+[.)]\s*', '', line)
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_run_font(run, bold=True)
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                set_run_font(run, italic=True)
            else:
                run = p.add_run(part)
                set_run_font(run)
        p.style = 'List Bullet' if line.startswith('- ') else 'List Number'
        set_para_spacing(p, before=0, after=3, line=18)
        i += 1
        continue

    # ── meta / caption lines ──────────────────────────────────────────────
    if line.startswith('**Fig.') or line.startswith('**Table'):
        add_caption(doc, line.replace('**', ''))
        i += 1
        continue

    # ── word count note ───────────────────────────────────────────────────
    if line.startswith('*Word count'):
        p = doc.add_paragraph()
        run = p.add_run(line.strip('*'))
        set_run_font(run, size=9, italic=True)
        set_para_spacing(p, before=12, after=0, line=14)
        i += 1
        continue

    # ── normal body paragraph ─────────────────────────────────────────────
    add_body(doc, line)
    i += 1

doc.save(DOCX)
print(f"✓ Saved {DOCX}")
print(f"  File size: {DOCX.stat().st_size / 1024:.0f} KB")
